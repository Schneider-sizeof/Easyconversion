import os
import fitz  # PyMuPDF
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image

# Define the Flask application
app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'your_secret_key'  # Required for session-based flash messages
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Route for the home page (upload form)
@app.route('/')
def index():
    return render_template('index.html')  # Make sure this HTML file exists in the templates folder

# Route for handling file conversion
@app.route('/convert', methods=['POST'])
def convert_file():
    uploaded_file = request.files['file']
    output_format = request.form['format']

    if uploaded_file:
        filename = secure_filename(uploaded_file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        uploaded_file.save(file_path)

        # Debugging logs (optional)
        print(f"Uploaded file: {filename}")
        print(f"Requested conversion to: {output_format}")

        # Check if the file is the same format as the requested output
        if filename.endswith(output_format):
            flash(f"No conversion needed: {output_format.upper()} format is the same as the uploaded file.", "error")
            return redirect(url_for('index'))

        # DOCX to PDF Conversion
        if filename.endswith('.docx') and output_format == 'pdf':
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename.rsplit('.', 1)[0] + '.pdf')
            convert_docx_to_pdf(file_path, pdf_path)
            return send_file(pdf_path, as_attachment=True)

        # PDF to DOCX Conversion
        elif filename.endswith('.pdf') and output_format == 'docx':
            docx_path = os.path.join(app.config['UPLOAD_FOLDER'], filename.rsplit('.', 1)[0] + '.docx')
            convert_pdf_to_docx(file_path, docx_path)
            return send_file(docx_path, as_attachment=True)

        # Image (PNG or JPG) to PDF Conversion
        elif filename.endswith(('.png', '.jpg', '.jpeg')) and output_format == 'pdf':
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename.rsplit('.', 1)[0] + '.pdf')
            convert_image_to_pdf(file_path, pdf_path)
            return send_file(pdf_path, as_attachment=True)

        # PDF to Image (PNG or JPG) Conversion
        elif filename.endswith('.pdf') and output_format in ('png', 'jpg'):
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename.rsplit('.', 1)[0] + f'.{output_format}')
            convert_pdf_to_image(file_path, image_path, output_format)
            return send_file(image_path, as_attachment=True)

        # If none of the above, just return the uploaded file
        return send_file(file_path, as_attachment=True)

    return redirect(url_for('index'))  # Redirect if no file is uploaded

# Conversion Functions
def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        doc = Document(docx_path)
        pdf = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        y = height - 40
        for paragraph in doc.paragraphs:
            pdf.drawString(40, y, paragraph.text)
            y -= 15
            if y < 40:
                pdf.showPage()
                y = height - 40
        pdf.save()
    except Exception as e:
        print(f"Error during DOCX to PDF conversion: {e}")

def convert_pdf_to_docx(pdf_path, docx_path):
    try:
        doc = Document()
        pdf = fitz.open(pdf_path)
        for page_num in range(pdf.page_count):
            page = pdf.load_page(page_num)
            text = page.get_text("text")
            doc.add_paragraph(text)
        doc.save(docx_path)
    except Exception as e:
        print(f"Error during PDF to DOCX conversion: {e}")

def convert_image_to_pdf(image_path, pdf_path):
    try:
        image = Image.open(image_path)
        image = image.convert('RGB')
        image.save(pdf_path, 'PDF', resolution=100.0)
    except Exception as e:
        print(f"Error during Image to PDF conversion: {e}")

def convert_pdf_to_image(pdf_path, image_path, image_format):
    try:
        pdf = fitz.open(pdf_path)
        page = pdf.load_page(0)  # Get the first page (you can modify this to get more pages)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img.save(image_path, format=image_format.upper())
    except Exception as e:
        print(f"Error during PDF to Image conversion: {e}")

# Ensure the app listens on the correct host and port
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))  # Get the port from environment variables or default to 5000
    app.run(host="0.0.0.0", port=port)  # Listen on all interfaces
